#!/usr/bin/env python3
"""
resume_submit.md → resume_submit.docx (제출용 이력서).

    python3 build_submit.py     # 먼저 resume.md → resume_submit.md
    python3 build_docx.py       # 그다음 docx

── 원본 규율 ─────────────────────────────────────────────────────────
원본은 resume.md 다. **resume.md 를 직접 읽지 않는다** — 상단 작업 노트와 하단 작성
메모가 그대로 딸려 들어가기 때문에, build_submit.py 가 그것을 걷어낸 resume_submit.md
를 입력으로 쓴다.

서식은 references/이력서_김태종_v1.1.0.docx 를 **템플릿으로 재사용** 한다. 새로
만들지 않고 그 파일의 본문만 비운 뒤 채우므로, 폰트·여백·스타일 정의가 그대로 유지된다.

산출물은 연봉·전화번호·주소가 들어간 전체판이다. .gitignore 대상 — 저장소에 올리지 않는다.

── 위계 표현 (2026.08.12 결정) ────────────────────────────────────────
참고 docx 는 들여쓰기 없이 평평했으나, 지금 이력서는 문제/방법/결과 3단 구조라
그대로 두면 무엇이 무엇에 속하는지 읽히지 않는다. 들여쓰기로만 위계를 준다:

    0단계  - **수행업무 제목**        굵게 · 들여쓰기 없음 · 위 여백
    1단계    - 기간/문제/방법/결과     0.25" · 라벨만 굵게
    2단계      - 세부 내용            0.50" · 앞에 · 글머리
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).parent
SRC = ROOT / "resume_submit.md"
TEMPLATE = ROOT / "references" / "이력서_김태종_v1.1.0.docx"
OUT = ROOT / "resume_submit.docx"

BODY = "normal"          # 템플릿이 쓰는 본문 스타일 이름 (소문자다)
INDENT = (0.0, 0.25, 0.5)

# 산출물에 반드시 남아야 하는 것 — 통째로 날아가면 잡는다
REQUIRED = ("개인신상", "경력기술서", "핵심 기술 요약")

# 작업 메모가 딸려 오면 중단한다 (build_submit.py 가 걸렀어야 하는 것)
LEAKS = ("※ 작성 메모", "<!--", "> 이 문서는")


def clear_body(doc: Document) -> None:
    """템플릿의 본문만 비운다. 스타일·섹션 설정은 남는다."""
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):      # 페이지 설정은 보존
            continue
        body.remove(child)


def add_runs(par, text: str) -> None:
    """**굵게** 와 `코드` 를 run 으로 나눠 넣는다."""
    for piece in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = "Menlo"
        else:
            par.add_run(piece)


def para(doc, text, level=0, bullet="", space_before=0):
    p = doc.add_paragraph(style=BODY)
    pf = p.paragraph_format
    pf.left_indent = Inches(INDENT[level])
    if space_before:
        pf.space_before = Pt(space_before)
    if bullet:
        p.add_run(bullet + " ")
    add_runs(p, text)
    return p


def build() -> Document:
    if not SRC.exists():
        sys.exit(f"원본이 없습니다: {SRC}\n먼저 `python3 build_submit.py` 를 실행하세요.")
    if not TEMPLATE.exists():
        sys.exit(f"서식 템플릿이 없습니다: {TEMPLATE}")

    text = SRC.read_text(encoding="utf-8")
    leaked = [m for m in LEAKS if m in text]
    if leaked:
        sys.exit(f"중단: 작업 메모가 남아 있습니다 → {leaked}\n"
                 f"`python3 build_submit.py` 를 다시 실행하세요. (docx 미생성)")

    doc = Document(str(TEMPLATE))
    clear_body(doc)

    for raw in text.split("\n"):
        line = raw.rstrip()
        s = line.strip()

        if not s or s == "---":
            continue

        m = re.match(r"^(#{1,3}) (.+)$", s)
        if m:
            n, title = len(m.group(1)), m.group(2).strip()
            if n == 1:
                doc.add_paragraph(title, style="Title")
            else:
                doc.add_paragraph(title, style=f"Heading {n}")
            continue

        # **<주요 수행업무>** 같은 소제목
        if s.startswith("**") and s.endswith("**"):
            para(doc, s, space_before=10)
            continue

        m = re.match(r"^( *)- (.*)$", line)
        if m:
            level = min(len(m.group(1)) // 2, 2)
            body = m.group(2).strip()
            if level == 0:
                para(doc, body, level=0, space_before=10)
            elif level == 1:
                para(doc, body, level=1)
            else:
                para(doc, body, level=2, bullet="·")
            continue

        para(doc, s)

    return doc


def main() -> None:
    doc = build()

    # 검사를 먼저 한다. 통과하지 못하면 파일을 쓰지 않는다.
    got = "\n".join(p.text for p in doc.paragraphs)
    missing = [r for r in REQUIRED if r not in got]
    if missing:
        sys.exit(f"중단: 필수 섹션이 없습니다 → {missing} (docx 미생성)")

    doc.save(str(OUT))
    print(f"생성 완료: {OUT.name} ({OUT.stat().st_size:,} bytes · 문단 {len(doc.paragraphs)}개)")
    print("주의: 연봉·전화번호·주소가 포함된 제출용 전체판입니다. 저장소에 올리지 마세요.")


if __name__ == "__main__":
    main()
